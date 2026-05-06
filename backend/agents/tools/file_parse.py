"""文件智能解析工具 — 自动识别格式，提取文本/表格/结构化内容

系统级工具：所有 agent 均可使用，无需单独配置。
支持格式：PDF、DOCX、XLSX、XLS、CSV、TXT、JSON、YAML、MD、图片（PNG/JPG/GIF/BMP/WebP/TIFF）等。
"""
import csv
import os

from agents.tool_registry import register_tool, ToolContext
from agents.workspace import safe_path

IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp", ".tiff", ".tif"}


@register_tool(read_only=True)
def file_parse(path: str, max_pages: int = 50, max_rows: int = 200, ctx: ToolContext = None) -> dict:
    """智能解析工作区内的文件，自动识别格式并提取文本和表格内容。

    支持格式：PDF、DOCX、XLSX、XLS、CSV、TXT、JSON、YAML、MD、图片（PNG/JPG/GIF/BMP/WebP/TIFF）。
    path 为工作区内相对路径（如 "inbox/流水.xlsx"），不是绝对路径。

    使用场景：
    - 用户上传了银行流水文件 → 解析后分析列结构、提取数据
    - 用户上传了报表模板 → 解析后识别占位符和结构
    - 用户上传了图片 → 返回图片元信息和 base64 数据

    参数说明：
    - path: 必需，文件在工作区中的相对路径
    - max_pages: PDF 最大解析页数，默认 50
    - max_rows: Excel/CSV 最大解析行数，默认 200

    返回格式：{"ok": true, "format": "excel|pdf|csv|...", "file": "文件名", "content": "解析内容文本", "total_chars": 1234}
    失败时返回：{"ok": false, "error": "错误原因", "format": "文件格式"}
    """
    abs_path = safe_path(ctx.agent_code, path)
    if not abs_path or not os.path.isfile(abs_path):
        return {"ok": False, "error": f"文件不存在: {path}"}

    ext = os.path.splitext(abs_path)[1].lower()

    if ext == ".pdf":
        return _parse_pdf(abs_path, max_pages)
    if ext == ".docx":
        return _parse_docx(abs_path)
    if ext in (".xlsx", ".xls"):
        return _parse_excel(abs_path, max_rows)
    if ext == ".csv":
        return _parse_csv(abs_path, max_rows)
    if ext in IMAGE_EXTS:
        return _parse_image(abs_path)
    return _parse_text(abs_path)


def _parse_pdf(filepath: str, max_pages: int) -> dict:
    try:
        import pdfplumber
    except ImportError:
        return {"ok": False, "error": "pdfplumber 未安装，无法解析 PDF"}

    pages = []
    try:
        with pdfplumber.open(filepath) as pdf:
            for i, page in enumerate(pdf.pages):
                if i >= max_pages:
                    pages.append(f"\n--- [已截断，共 {len(pdf.pages)} 页，仅展示前 {max_pages} 页] ---")
                    break
                text = page.extract_text() or ""
                if text.strip():
                    pages.append(f"--- 第 {i + 1} 页 ---\n{text}")

                tables = page.extract_tables()
                for ti, table in enumerate(tables):
                    if not table:
                        continue
                    header = table[0] if table else []
                    rows = table[1:] if len(table) > 1 else []
                    table_text = _format_table(header, rows)
                    pages.append(f"[表格 {ti + 1}]\n{table_text}")

        content = "\n\n".join(pages)
        if not content.strip():
            return {"ok": False, "error": "PDF 内容为空（可能是扫描件/图片 PDF，需要 OCR 支持）",
                    "format": "pdf", "file": os.path.basename(filepath)}

        return {
            "ok": True,
            "format": "pdf",
            "file": os.path.basename(filepath),
            "content": content,
            "total_chars": len(content),
        }
    except Exception as e:
        return {"ok": False, "error": f"PDF 解析失败: {e}", "format": "pdf"}


def _parse_docx(filepath: str) -> dict:
    try:
        from docx import Document
    except ImportError:
        return {"ok": False, "error": "python-docx 未安装，无法解析 DOCX"}

    try:
        doc = Document(filepath)
        parts = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        for table in doc.tables:
            header = [cell.text.strip() for cell in table.rows[0].cells] if table.rows else []
            rows = []
            for row in table.rows[1:]:
                rows.append([cell.text.strip() for cell in row.cells])
            if header:
                parts.append(_format_table(header, rows))

        content = "\n\n".join(parts)
        return {
            "ok": True,
            "format": "docx",
            "file": os.path.basename(filepath),
            "content": content,
            "total_chars": len(content),
        }
    except Exception as e:
        return {"ok": False, "error": f"DOCX 解析失败: {e}", "format": "docx"}


def _parse_excel(filepath: str, max_rows: int) -> dict:
    from core.excel_compat import load_workbook_compat

    try:
        wb = load_workbook_compat(filepath, data_only=True)
        parts = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = []
            for i, row in enumerate(ws.iter_rows(values_only=True)):
                if i >= max_rows:
                    rows.append(["... (已截断)"])
                    break
                rows.append([str(c) if c is not None else "" for c in row])

            if rows:
                header = rows[0]
                data = rows[1:]
                parts.append(f"[Sheet: {sheet_name}]\n" + _format_table(header, data))

        wb.close()
        content = "\n\n".join(parts)
        return {
            "ok": True,
            "format": "excel",
            "file": os.path.basename(filepath),
            "content": content,
            "total_chars": len(content),
        }
    except Exception as e:
        return {"ok": False, "error": f"Excel 解析失败: {e}", "format": "excel"}


def _parse_csv(filepath: str, max_rows: int) -> dict:
    try:
        with open(filepath, "r", encoding="utf-8-sig", errors="replace") as f:
            reader = csv.reader(f)
            rows = []
            for i, row in enumerate(reader):
                if i >= max_rows:
                    rows.append(["... (已截断)"])
                    break
                rows.append(row)

        if not rows:
            return {"ok": False, "error": "CSV 文件为空", "format": "csv"}

        header = rows[0]
        data = rows[1:]
        content = _format_table(header, data)

        return {
            "ok": True,
            "format": "csv",
            "file": os.path.basename(filepath),
            "content": content,
            "total_chars": len(content),
        }
    except Exception as e:
        return {"ok": False, "error": f"CSV 解析失败: {e}", "format": "csv"}


def _parse_text(filepath: str) -> dict:
    try:
        with open(filepath, "r", encoding="utf-8", errors="replace") as f:
            content = f.read()
        return {
            "ok": True,
            "format": "text",
            "file": os.path.basename(filepath),
            "content": content[:50000],
            "total_chars": len(content),
            "truncated": len(content) > 50000,
        }
    except Exception as e:
        return {"ok": False, "error": f"文件读取失败: {e}"}


def _parse_image(filepath: str) -> dict:
    """解析图片文件，返回基本信息和 base64 数据"""
    import base64
    import mimetypes

    try:
        file_size = os.path.getsize(filepath)
        mime_type = mimetypes.guess_type(filepath)[0] or "image/png"

        width, height = None, None
        try:
            from PIL import Image
            with Image.open(filepath) as img:
                width, height = img.size
        except ImportError:
            pass

        # 限制图片大小，超过 5MB 不传 base64
        b64 = None
        data_uri = None
        if file_size <= 5 * 1024 * 1024:
            with open(filepath, "rb") as f:
                img_data = f.read()
            b64 = base64.b64encode(img_data).decode("ascii")
            data_uri = f"data:{mime_type};base64,{b64}"

        info_parts = [
            f"文件: {os.path.basename(filepath)}",
            f"类型: {mime_type}",
            f"大小: {file_size / 1024:.1f}KB",
            "",
            "请仔细分析这张图片的内容，识别其中的文字、表格、数据、界面元素等所有信息。",
            "根据图片内容和用户的原始需求，立即执行后续操作（如调整配置、创建模板、修改规则等）。",
        ]
        if width and height:
            info_parts.append(f"尺寸: {width}x{height}")

        return {
            "ok": True,
            "format": "image",
            "file": os.path.basename(filepath),
            "mime_type": mime_type,
            "width": width,
            "height": height,
            "file_size_kb": round(file_size / 1024, 1),
            "content": "\n".join(info_parts),
            "data_uri": data_uri,
        }
    except Exception as e:
        return {"ok": False, "error": f"图片解析失败: {e}", "format": "image"}


def _format_table(header: list[str], rows: list[list[str]]) -> str:
    if not header:
        return ""
    col_widths = [len(str(h)) for h in header]
    for row in rows:
        for i, cell in enumerate(row):
            if i < len(col_widths):
                col_widths[i] = max(col_widths[i], len(str(cell)))

    def fmt_row(cells):
        parts = []
        for i, c in enumerate(cells):
            w = col_widths[i] if i < len(col_widths) else 10
            parts.append(str(c).ljust(w))
        return " | ".join(parts)

    lines = [fmt_row(header)]
    lines.append("-+-".join("-" * w for w in col_widths))
    for row in rows[:100]:
        lines.append(fmt_row(row))
    if len(rows) > 100:
        lines.append(f"... 共 {len(rows)} 行，仅展示前 100 行")

    return "\n".join(lines)
